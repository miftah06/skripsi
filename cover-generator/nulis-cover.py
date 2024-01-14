import os
import csv
from fpdf import FPDF
from docx import Document

def bootstrap():
    os.system("python bootstrap.py")

def builder():
    os.system("python builder.py")

def beauty_pdf(data):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Setting Font
    pdf.set_font("Arial", size=12)

    # Judul
    pdf.set_font("Arial", style='B', size=18)
    pdf.cell(0, 10, data['Judul'][0], ln=True, align='C')
    pdf.ln(10)

    # Logo
    if data['Logo'][0]:
        pdf.image(data['Logo'][0], x=50, w=100)
        pdf.ln(40)
    else:
        pdf.ln(10)

    # Teks
    pdf.set_font("Arial", style='', size=12)
    pdf.multi_cell(0, 10, data['Teks'][0])
    pdf.ln(10)

    # Oleh
    pdf.set_font("Arial", style='I', size=12)
    pdf.cell(0, 10, f"Oleh: {data['Oleh'][0]}", ln=True)
    pdf.ln(10)

    # Input Numerik
    pdf.set_font("Arial", style='', size=12)
    pdf.cell(0, 10, f"Numerik: {data['Input Numerik'][0]}", ln=True)
    pdf.ln(5)

    # Input 1
    pdf.cell(0, 10, f"Input 1: {data['Input 1'][0]}", ln=True)
    pdf.ln(5)

    # Input 2
    pdf.cell(0, 10, f"Input 2: {data['Input 2'][0]}", ln=True)
    pdf.ln(5)

    # Tahun
    pdf.cell(0, 10, f"Tahun: {data['Tahun'][0]}", ln=True)

    # Menyimpan PDF
    pdf.output('output-cover.pdf')

def buat_cover_doc(data):
    doc = Document()

    # Judul
    doc.add_heading(data['Judul'][0], level=1).alignment = 1  # 1 for center alignment

    # Teks
    doc.add_paragraph(data['Teks'][0])

    # Oleh
    doc.add_paragraph(f"Oleh: {data['Oleh'][0]}")

    # Input Numerik
    doc.add_paragraph(f"Numerik: {data['Input Numerik'][0]}")

    # Input 1
    doc.add_paragraph(f"Input 1: {data['Input 1'][0]}")

    # Input 2
    doc.add_paragraph(f"Input 2: {data['Input 2'][0]}")

    # Tahun
    doc.add_paragraph(f"Tahun: {data['Tahun'][0]}")

    # Menyimpan dokumen ke file .doc
    doc.save('output-cover.doc')

def buat_cover_pdf():
    print("Selamat datang di skrip cover.py!\n")

    # Meminta input dari pengguna
    judul = input("Masukkan judul: ").strip()
    logo = input("Masukkan logo (kosongkan jika tidak ada): ").strip()
    teks = input("Masukkan judul tugas (contoh: Skripsi): ").strip()
    nama = input("Masukkan nama: ").strip()
    numerik = input("Masukkan input numerik (contoh: NIM): ").strip()
    input_1 = input("Masukkan input 1 (contoh: Fakultas): ").strip()
    input_2 = input("Masukkan input 2 (contoh: Universitas): ").strip()
    tahun = input("Masukkan tahun (contoh: 2024): ").strip()

    # Membuat data
    data = {
        'Judul': [judul],
        'Logo': [logo],
        'Teks': [teks],
        'Oleh': [nama],
        'Input Numerik': [numerik],
        'Input 1': [input_1],
        'Input 2': [input_2],
        'Tahun': [tahun]
    }

    # Membuat PDF dengan estetika
    beauty_pdf(data)

    print("\nProses selesai. File PDF yang indah tersedia di output-cover.pdf.")

    # Membuat dokumen Word
    buat_cover_doc(data)
    print("File Doc yang indah telah berhasil dibuat dalam file output-cover.doc.")

if __name__ == "__main__":
    print("Selamat datang di skrip cover.py!\n")
    
    print("1. Menjalankan skrip Bootstrap untuk membuat file CSV.")
    bootstrap()

    print("\n2. Menjalankan skrip Builder untuk membuat file HTML.")
    builder()

    # Membaca data dari file CSV
    with open('output.csv', newline='', encoding='utf-8') as csvfile:
        reader = csv.DictReader(csvfile)
        data = {key: [] for key in reader.fieldnames}

        for row in reader:
            for key in reader.fieldnames:
                data[key].append(row[key].strip() if row[key] else '')

    print("\n3. Menjalankan skrip Beauty-PDF untuk membuat PDF yang indah.")
    # Memanggil fungsi beauty_pdf dengan menyediakan data yang dibutuhkan
    beauty_pdf(data)

    print("\nProses selesai. File PDF yang indah tersedia di output-cover.pdf.")

    # Meminta input dari pengguna
    judul = input("Masukkan judul: ").strip()
    logo = input("Masukkan logo: ").strip()
    teks = input("Masukkan judul tugas (contoh: Skripsi): ").strip()
    nama = input("Masukkan nama: ").strip()
    numerik = input("Masukkan input numerik (contoh: NIM): ").strip()
    input_1 = input("Masukkan input 1 (contoh: Fakultas): ").strip()
    input_2 = input("Masukkan input 2 (contoh: Universitas): ").strip()
    tahun = input("Masukkan tahun (contoh: 2024): ").strip()
    
    # Membuat dokumen Word
    buat_cover_doc(data)
    print("File Doc yang indah telah berhasil dibuat dalam file output-cover.doc.")
